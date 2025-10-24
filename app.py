import streamlit as st
import pandas as pd
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from datetime import datetime
import io


# --- Helper Functions to match original formatting ---

def format_time_condensed(time_str):
    """Formats time to H:MM without AM/PM and leading zeros."""
    if pd.isna(time_str) or time_str == '':
        return ''
    try:
        time_obj = datetime.strptime(str(int(float(time_str))).zfill(4), '%H%M')
        formatted_time = time_obj.strftime('%I:%M')
        if formatted_time.startswith('0'):
            return formatted_time[1:]
        return formatted_time
    except (ValueError, TypeError):
        return ''


def abbreviate_title(title):
    """Shortens course titles to match the original document's format."""
    if pd.isna(title):
        return ''
    title = str(title)
    # Comprehensive dictionary of replacements
    replacements = {
        'Anatomy & Physiology': 'A & P',
        'Bioenergetics and Systems': 'Bioenergetics',
        'Genomes and Evolution': 'Genome Evol',
        'Medical Microbiology': 'Med Micro',
        'Earth/Life Sci for Educators': 'Life Sci Ed',
        'Biostatistics': 'Biostats',
        'Biology Capstone Seminar': 'Capstone',
        'Insect Biology': 'Insect Bio',
        'Science in the Public Domain': 'SCI Pub Dom',
        'Ecological Community:San Diego': 'Ecol Comm',
        'Research Methods': 'Res Meth',
        'Cell Physiology': 'Cell Phys',
        'Vertebrate Physiology': 'Vert Phys',
        'Microbiology': 'Micro',
        'Research Project': 'Res Proj',
        'Techniques: Molecular Biology': 'Molec Tech',
        'Comp. Anat. of Vertebrates': 'Comp An Vert',
        'Comparative Anatomy (Linked with Human Evolution)': 'Comp Ant (linked)',
        'Invertebrate Zoology': 'Invert Zoo',
        'Peoples, Plagues and Microbes': 'Ppl Plag Micro',
        'Ecol Evol Infectious Disease': 'EEID',
        'Life Changing Biology': 'Life Change Bio',
        'Immunology': 'Immuno',
        'Laboratory': '',
        'Lab': ''
    }
    # Iterate through keys sorted by length (longest first) to handle overlaps
    for old in sorted(replacements, key=len, reverse=True):
        new = replacements[old]
        title = title.replace(old, new)
    return title


def correct_instructor_name(name):
    """Corrects specific instructor names to match the original doc."""
    if pd.isna(name): return ''
    name = str(name).strip().upper()
    corrections = {
        'NYHOLT DE PRADA': 'PRADA',
        'RECART GONZALEZ': 'RECART-GONZALEZ',
        'FLEMING-DAVIES': 'FLEMING-DAVIES'  # Ensures consistent formatting
    }
    return corrections.get(name, name)


# --- Core Logic from your script (adapted for new formatting) ---

def load_schedule_data(uploaded_file):
    """Loads data from an uploaded CSV or Excel file."""
    try:
        if uploaded_file.name.endswith('.csv'):
            df = pd.read_csv(uploaded_file)
        elif uploaded_file.name.endswith('.xlsx') or uploaded_file.name.endswith('.xls'):
            df = pd.read_excel(uploaded_file, engine='openpyxl')
        else:
            st.error("Unsupported file format. Please upload a CSV or Excel file.")
            return None
        return df
    except Exception as e:
        st.error(f"Error reading the file: {e}")
        return None


def parse_time(time_str):
    if pd.isna(time_str) or time_str == '': return None
    try:
        time_str = str(time_str).strip().zfill(4)
        hours = int(time_str[:-2])
        minutes = int(time_str[-2:])
        return hours * 60 + minutes
    except:
        return None


def get_day_of_week(row):
    days = []
    for day_char, day_full in [('M', 'Monday'), ('T', 'Tuesday'), ('W', 'Wednesday'), ('R', 'Thursday'),
                               ('F', 'Friday')]:
        # Check if the column exists and if its value matches the day_char
        if day_char in row and str(row.get(day_char)).strip() == day_char:
            days.append(day_full)
    return days


def process_schedule_data(df):
    """
    Processes the loaded DataFrame into a room schedule dictionary.
    Assumes column validation has already happened.
    """
    room_schedule = {}
    
    # Drop rows where essential data for scheduling is missing
    df_cleaned = df.dropna(subset=['BLDG', 'ROOM', 'BEGIN', 'END'])

    # The .copy() prevents a SettingWithCopyWarning
    df_cleaned = df_cleaned.copy() 
    
    # Strip whitespace from key columns
    strip_cols = ['BLDG', 'ROOM', 'SUBJ', 'CRSE #', 'LAST NAME', 'M', 'T', 'W', 'R', 'F']
    for col in strip_cols:
        if col in df_cleaned.columns:
            # Ensure data is string type before stripping
            df_cleaned[col] = df_cleaned[col].astype(str).str.strip()

    df_cleaned['Course'] = df_cleaned['SUBJ'] + df_cleaned['CRSE #']
    df_cleaned['Instructor'] = df_cleaned['LAST NAME'].apply(correct_instructor_name)
    df_cleaned['Days'] = df_cleaned.apply(get_day_of_week, axis=1)

    for _, row in df_cleaned.iterrows():
        try:
            room_name = f"{row['BLDG'].replace('SCST', 'ST')}{int(float(row['ROOM']))}"
        except (ValueError, TypeError):
            continue
        
        begin_time = parse_time(row['BEGIN'])
        if begin_time is None: continue

        for day in row['Days']:
            if day not in room_schedule: room_schedule[day] = {}
            if room_name not in room_schedule[day]: room_schedule[day][room_name] = []

            room_schedule[day][room_name].append({
                'Begin': row['BEGIN'],
                'End': row['END'],
                'Course': row['Course'],
                'Title': row['TITLE'],
                'Instructor': row['Instructor'],
                'BeginMinutes': begin_time,
                'IsMorning': begin_time < 720
            })

    for day in room_schedule:
        for room in room_schedule[day]:
            room_schedule[day][room].sort(key=lambda x: x['BeginMinutes'])
    return room_schedule


def create_room_use_chart(room_schedule):
    doc = Document()
    # Set to Landscape
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    new_width, new_height = section.page_height, section.page_width
    section.page_width = new_width
    section.page_height = new_height
    # Set margins
    for sec in doc.sections:
        sec.top_margin = Inches(0.5)
        sec.bottom_margin = Inches(0.5)
        sec.left_margin = Inches(0.5)
        sec.right_margin = Inches(0.5)

    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run('Room Use Chart for the Biology Laboratories')
    font_title = run_title.font
    font_title.name = 'Times New Roman'
    font_title.size = Pt(20)
    font_title.bold = True

    # Table
    days_of_week = ['Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday']
    all_rooms = ['ST225', 'ST227', 'ST229', 'ST242', 'ST325', 'ST327', 'ST330', 'ST429']

    table = doc.add_table(rows=1, cols=len(all_rooms) + 1)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].width = Inches(1.0)

    # Table Header Content
    p_hdr_legend = hdr_cells[0].paragraphs[0]
    p_hdr_legend.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_b_hdr = p_hdr_legend.add_run('B=Morning\n')
    font_b_hdr = run_b_hdr.font
    font_b_hdr.name = 'Times New Roman'
    font_b_hdr.size = Pt(9)
    font_b_hdr.bold = True
    font_b_hdr.color.rgb = RGBColor(0, 0, 255)
    run_g_hdr = p_hdr_legend.add_run('G=Afternoon')
    font_g_hdr = run_g_hdr.font
    font_g_hdr.name = 'Times New Roman'
    font_g_hdr.size = Pt(9)
    font_g_hdr.bold = True
    font_g_hdr.color.rgb = RGBColor(0, 128, 0)

    for i, col_name in enumerate(all_rooms, 1):
        p_hdr = hdr_cells[i].paragraphs[0]
        p_hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_hdr.add_run(col_name)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(20)
        run.font.bold = True
        hdr_cells[i].width = Inches(1.25)

    # Table Body
    for day in days_of_week:
        row_cells = table.add_row().cells
        p_day = row_cells[0].paragraphs[0]
        p_day.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_day = p_day.add_run(day[:3])
        run_day.font.name = 'Times New Roman'
        run_day.font.size = Pt(20)
        run_day.bold = True

        for j, room_name in enumerate(all_rooms, 1):
            val = room_schedule.get(day, {}).get(room_name, [])
            para = row_cells[j].paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Set vertical alignment
            morning = [v for v in val if v['IsMorning']]
            afternoon = [v for v in val if not v['IsMorning']]
            if len(morning) == 0 and len(afternoon) == 1:
                row_cells[j].vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
            elif len(afternoon) == 0 and len(morning) == 1:
                row_cells[j].vertical_alignment = WD_ALIGN_VERTICAL.TOP
            else:
                row_cells[j].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            if val:
                for idx, v in enumerate(val):
                    if idx > 0: para.add_run("\n\n")

                    begin = format_time_condensed(v['Begin'])
                    end = format_time_condensed(v['End'])
                    title = abbreviate_title(v['Title'])
                    text = f"{begin}-{end}\n{v['Course']}\n{title}\n{v['Instructor']}"

                    run = para.add_run(text)
                    font = run.font
                    font.name = 'Times New Roman'
                    font.bold = True
                    font.size = Pt(9)
                    font.color.rgb = RGBColor(0, 0, 255) if v['IsMorning'] else RGBColor(0, 128, 0)
    return doc


# --- Streamlit App UI ---
st.set_page_config(page_title="Bio Room Use Chart Generator", layout="wide")
st.title(":memo: Bio Room Use Chart Generator for Dept Chair")
st.write("This tool converts a class schedule into a Room Use Chart.")

# Define the original required columns based on the original template
REQUIRED_COLUMNS = [
    'SUBJ', 'CRSE #', 'SEC #', 'TITLE', 'ATTRIBUTE', 'UNITS', 'M', 'T', 'W', 'R', 'F', 
    'BEGIN', 'END', 'BLDG', 'ROOM', 'ENROLLMENT', 'LAST NAME', 'FIRST NAME'
]

st.header("Upload Your Schedule File")
uploaded_file = st.file_uploader("", type=['csv', 'xlsx', 'xls'], label_visibility="collapsed")

# Initialize session state variables
if 'df_loaded' not in st.session_state:
    st.session_state.df_loaded = None
if 'file_valid' not in st.session_state:
    st.session_state.file_valid = False
if 'chart_data' not in st.session_state:
    st.session_state.chart_data = None
if 'last_uploaded_filename' not in st.session_state:
    st.session_state.last_uploaded_filename = None

if uploaded_file is not None:
    # Check if a new file has been uploaded
    if uploaded_file.name != st.session_state.last_uploaded_filename:
        st.session_state.last_uploaded_filename = uploaded_file.name
        st.session_state.df_loaded = load_schedule_data(uploaded_file)
        st.session_state.file_valid = False # Reset validation
        st.session_state.chart_data = None # Reset generated chart
        
        if st.session_state.df_loaded is not None:
            # Validate the columns
            missing_cols = [col for col in REQUIRED_COLUMNS if col not in st.session_state.df_loaded.columns]
            if missing_cols:
                st.error(f"The uploaded file is missing the following required columns: {', '.join(missing_cols)}")
                st.session_state.file_valid = False
            else:
                st.success("File processed successfully! Click the button below to generate your chart.")
                st.session_state.file_valid = True

# Show the "Generate" button only if the file is valid and chart isn't generated
if st.session_state.file_valid and st.session_state.chart_data is None:
    if st.button("Generate Room Chart"):
        with st.spinner('Generating your chart...'):
            room_schedule = process_schedule_data(st.session_state.df_loaded)
            if room_schedule:
                doc = create_room_use_chart(room_schedule)
                bio = io.BytesO()
                doc.save(bio)
                st.session_state.chart_data = bio.getvalue()
                st.success("Chart generated!")
            else:
                st.warning(
                    "Could not generate a chart. Please check that your file contains the correct data.")
                st.session_state.chart_data = None

# Show the "Download" button only if the chart data exists
if st.session_state.chart_data is not None:
    st.download_button(
        label="Download Word Document",
        data=st.session_state.chart_data,
        file_name=f"Room_Use_Chart_{datetime.now().strftime('%Y%m%d')}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

st.markdown("---")
st.header("How to Use This App")

# Updated template CSV string to match the original template
template_csv = "SUBJ,CRSE #,SEC #,TITLE,ATTRIBUTE,UNITS,M,T,W,R,F,BEGIN,END,BLDG,ROOM,ENROLLMENT,LAST NAME,FIRST NAME\n"

st.markdown("**Step 1: Get the Template**\n- Click the button below to download the required template.")
st.download_button(
    label="Download Template CSV", 
    data=template_csv, 
    file_name="Template_Schedule.csv", 
    mime="text/csv"
)
st.markdown(f"""
**Step 2: Prepare Your Data**
- Open the downloaded template (`Room_Chart_Template.csv`) in Excel.
- :exclamation:**Crucial:** Copy your class schedule data into the appropriate columns. Do NOT change any column headers.:exclamation:
- Save the file as a .xlsx or .csv file. 

**Step 3: Upload Your File**
- Use the uploader at the top of the page to upload your file.
- If you get an error message, ensure that your file contains all the data required in the Room_Chart_Template.csv. 

**Step 4: Generate and Download**
- Click the **"Generate Room Chart"** button.
- Once generated, click the **"Download Word Document"** button to download the room chart.
""")

